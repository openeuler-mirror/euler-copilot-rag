import os
import re
from docx import Document
from scripts.corpus.handler.document_handler import DocumentHandler
from logger import get_logger
from multiprocessing import Process

logger = get_logger()


class CorpusHandler():
    @staticmethod
    def get_all_file_paths(directory):
        file_paths = []
        for dirpath, dirnames, filenames in os.walk(directory):
            for filename in filenames:
                full_path = os.path.join(dirpath, filename)
                file_paths.append(full_path)
        return file_paths

    @staticmethod
    def clean_string(s):
        return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', s)

    @staticmethod
    def write_para_to_docx(language, tar_dir, para_cnt, file_name, para_list):
        for para in para_list:
            try:
                name = os.path.splitext(file_name)[0]
                document = Document()
                document.add_paragraph('<'+name+'>'+'\n')
                para = CorpusHandler.clean_string(para)
                document.add_paragraph(para)
                document.save(os.path.join(tar_dir, name+'_paragraph'+str(para_cnt)+'.docx'))
                para_cnt += 1
            except Exception as e:
                if language == 'zh':
                    logger.error(f'片段写入失败由于{e}')
                else:
                    logger.error(f'Fragment write failed due to {e}')
                print(para)

    @staticmethod
    def change_document_to_para(language, src_dir, tar_dir, para_chunk=1024, num_cores=8):
        all_files = CorpusHandler.get_all_file_paths(src_dir)
        file_to_para_dict = {}
        for dir in all_files:
            try:
                para_list = DocumentHandler.get_content_list_from_file(dir, para_chunk, num_cores)
            except Exception as e:
                if language == 'zh':
                    logger.error(f'文件 {src_dir}转换为片段失败，由于错误{e}')
                else:
                    logger.error(f'Conversion of file {src_dir} to fragment failed due to error {e}')
                continue
            if len(para_list) == 0:
                continue
            file_name = os.path.basename(dir)
            if len(para_list) != 0:
                file_to_para_dict[dir] = para_list
            if os.cpu_count() is None:
                return []
            num_cores = min(num_cores, os.cpu_count()//2)
            num_cores = min(8, min(num_cores, len(para_list)))
            num_cores = max(num_cores, 1)
            task_chunk = len(para_list)//num_cores
            processes = []
            for i in range(num_cores):
                st = i*task_chunk
                en = (i+1)*task_chunk
                if i == num_cores-1:
                    en = len(para_list)
                p = Process(target=CorpusHandler.write_para_to_docx, args=(
                    language,tar_dir, i*task_chunk, file_name, para_list[st:en]))
                processes.append(p)
                p.start()
            for i in range(len(processes)):
                processes[i].join()
        return file_to_para_dict
