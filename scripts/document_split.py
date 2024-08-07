import os
import nltk
from nltk.tokenize import word_tokenize
from bs4 import BeautifulSoup
import PyPDF2
from docx import Document
import markdown
from multiprocessing import Process, Manager

nltk_data_path = "./nltk_data"

if os.path.exists(nltk_data_path):
    nltk.data.path.append(nltk_data_path)


def tokenize_text(text):
    tokens = word_tokenize(text)
    return tokens


def tokenize_html(html):
    soup = BeautifulSoup(html, 'html.parser')
    text = soup.get_text()
    tokens = word_tokenize(text)
    return tokens


def tokenize_pdf(pdf_path):
    tokens = []
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text = page.extract_text()
            tokens.extend(word_tokenize(text))
    return tokens


def extract_paragraph(paragraph):
    return word_tokenize(paragraph.text)+['\n']


def extract_table(table):
    tokens = []
    tokens.extend('\n\n')
    header_printed = False

    for row in table.rows:
        row_tokens = []
        if not header_printed:
            header_line = ['-' * (len(str(cell.text).strip()) + 2) for cell in row.cells]
            header_printed = True

        for cell in row.cells:
            row_tokens.extend(word_tokenize(str(cell.text).strip()))
            row_tokens.append(' | ')

        tokens.extend(row_tokens)
        tokens.append('\n')

    if header_printed:
        tokens.extend(['-'.join(header_line), '\n'])
    tokens.extend('\n\n')
    return tokens

def create_element_map(doc):
    element_map = {}
    for paragraph in doc.paragraphs:
        element_map[paragraph._element] = paragraph
    for table in doc.tables:
        element_map[table._element] = table
    return element_map

def get_sub_tokens(pid, element_map, sub_el_list, q):
    for element in sub_el_list:
        sub_tokens = []
        if element in element_map:
            obj = element_map[element]
            if obj._element.tag.endswith('p'):
                sub_tokens = extract_paragraph(obj)
            elif obj._element.tag.endswith('tbl'):
                sub_tokens = extract_table(obj)
        q.put((pid, sub_tokens))
    
    q.put(None)


def tokenize_docx(docx_path):
    doc = Document(docx_path)
    tmp_tokens = []

    el_list = []
    for element in doc.element.body:
        el_list.append(element)
    num_cores = os.cpu_count()
    if num_cores is None:
        return []
    num_cores//=2
    num_cores=max(num_cores,1)
    num_cores = min(8,min(num_cores, len(el_list)))
    chunk_sz = len(el_list)//num_cores
    for i in range(num_cores):
        tmp_tokens.append([])
    processes = []
    q = Manager().Queue()
    element_map = create_element_map(doc)
    for i in range(num_cores):
        p = Process(target=get_sub_tokens, args=(i, element_map, el_list[chunk_sz*i:min(chunk_sz*(i+1), len(el_list))], q))
        processes.append(p)
        p.start()
    tokens = []
    task_finish_cnt = 0
    tmp_cnt=0
    while task_finish_cnt < num_cores:
        tmp = q.get()
        tmp_cnt+=1
        if tmp is None:
            task_finish_cnt += 1
            if task_finish_cnt == num_cores:
                break
            continue
        else:
            id, sub_tokens = tmp
        if tmp_cnt%1000==0:
            print(tmp_cnt)
        tmp_tokens[id].extend(sub_tokens)
    for sub_tokens in tmp_tokens:
        tokens.extend(sub_tokens)
    return tokens


def tokenize_md(md_path):
    with open(md_path, 'r', encoding='utf-8', errors="ignore") as file:
        md_text = file.read()
        html_text = markdown.markdown(md_text)
        tokens = tokenize_html(html_text)
    return tokens


def split_into_paragraphs(words, max_paragraph_length=1024):
    if max_paragraph_length == -1:
        return words
    paragraphs = []
    current_paragraph = ""
    for word in words:
        if len(current_paragraph) + len(word) <= max_paragraph_length:
            current_paragraph += word
        else:
            paragraphs.append(current_paragraph)
            current_paragraph = word

    if current_paragraph:
        paragraphs.append(current_paragraph.strip())
    return paragraphs


def get_paragraphs_from_file(file_path, max_paragraph_length=1024):
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors="ignore") as file:
            text = file.read()
            tokens = list(text)
            paragraphs = split_into_paragraphs(tokens, max_paragraph_length)
            return paragraphs
    elif file_extension == '.html':
        with open(file_path, 'r', encoding='utf-8', errors="ignore") as file:
            html_text = file.read()
            tokens = tokenize_html(html_text)
            paragraphs = split_into_paragraphs(tokens, max_paragraph_length)
            return paragraphs
    elif file_extension == '.pdf':
        tokens = tokenize_pdf(file_path)
        paragraphs = split_into_paragraphs(tokens, max_paragraph_length)
        return paragraphs
    elif file_extension == '.docx':
        tokens = tokenize_docx(file_path)
        paragraphs = split_into_paragraphs(tokens, max_paragraph_length)
        return paragraphs
    elif file_extension == '.md':
        tokens = tokenize_md(file_path)
        paragraphs = split_into_paragraphs(tokens, max_paragraph_length)
        return paragraphs
    else:
        return []
