import os
import re
from docx import Document
from document_split import get_paragraphs_from_file
from logger import get_logger
from multiprocessing import Process

logger = get_logger()


def get_all_file_paths(directory):
    file_paths = []  # 存储所有文件的绝对路径
    for dirpath, dirnames, filenames in os.walk(directory):
        for filename in filenames:
            full_path = os.path.join(dirpath, filename)  # 获取文件的绝对路径
            file_paths.append(full_path)
    return file_paths

def clean_string(s):
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', s)

def write_para_to_docx(tar_dir,para_cnt,file_name,para_list):
    for para in para_list:
            try:
                name = os.path.splitext(file_name)[0]
                document = Document()
                document.add_paragraph('<'+name+'>'+'\n')
                para=clean_string(para)
                document.add_paragraph(para)
                document.save(os.path.join(tar_dir, name+'_片段'+str(para_cnt)+'.docx'))
                para_cnt += 1
            except Exception as e:
                logger.error(f'片段写入失败由于{e}')
                print(para)
def change_document_to_para(src_dir, tar_dir, chunk=1024):
    all_files = get_all_file_paths(src_dir)
    file_name_list=[]
    for dir in all_files:
        try:
            para_list = get_paragraphs_from_file(dir, chunk)
        except Exception as e:
            logger.error(f'文件 {src_dir}转换为片段失败，由于错误{e}')
            continue
        file_name = os.path.basename(dir)
        if len(para_list)!=0:
            file_name_list.append(file_name)
        num_cores = os.cpu_count()
        if num_cores is None:
            return []
        num_cores//=2
        num_cores=max(num_cores,1)
        num_cores = min(8,min(num_cores, len(para_list)))
        chunk=len(para_list)//num_cores
        processes=[]
        for i in range(num_cores):
            p = Process(target=write_para_to_docx, args=(tar_dir,i*chunk, file_name, para_list[i*chunk:min(chunk*(i+1), len(para_list))]))
            processes.append(p)
            p.start()
        for i in range(len(processes)):
            processes[i].join() 
    return file_name_list
